from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_METRICS_DIR = PROJECT_ROOT / "artifacts" / "metrics"
DEFAULT_FIGURES_DIR = PROJECT_ROOT / "artifacts" / "figures"
DEFAULT_OUTPUT = PROJECT_ROOT / "Seminarski rad.docx"


def _load_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def _metric(payload: dict[str, Any], key: str, default: str = "nije pokrenuto") -> str:
    value = payload.get(key)
    if isinstance(value, float):
        return f"{value:.4f}"
    if value is None:
        return default
    return str(value)


def _add_metric_table(document: Document, title: str, rows: list[tuple[str, str]]) -> None:
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Metrika"
    header[1].text = "Vrednost"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _add_figure_if_exists(document: Document, figures_dir: Path, filename: str, caption: str) -> None:
    path = figures_dir / filename
    if not path.exists():
        return
    document.add_paragraph(caption)
    document.add_picture(str(path), width=Inches(5.8))


def generate_report(
    metrics_dir: str | Path = DEFAULT_METRICS_DIR,
    figures_dir: str | Path = DEFAULT_FIGURES_DIR,
    output_path: str | Path = DEFAULT_OUTPUT,
) -> Path:
    metrics_path = Path(metrics_dir)
    figures_path = Path(figures_dir)
    output = Path(output_path)

    classification = _load_json(metrics_path / "classification_metrics.json")
    regression = _load_json(metrics_path / "regression_metrics.json")
    encoder = _load_json(metrics_path / "encoder_metrics.json")
    prompt = _load_json(metrics_path / "prompt_metrics.json")
    finetuning = _load_json(metrics_path / "finetuning_metrics.json")
    error_analysis = _load_json(metrics_path / "error_analysis.json")

    document = Document()
    document.add_heading(
        "Implementacija encoder i generativnih modela masinskog ucenja na turistickim recenzijama",
        0,
    )
    document.add_paragraph("Drzavni univerzitet u Novom Pazaru")
    document.add_paragraph("Mentor: Ulfeta Marovac")
    document.add_paragraph("Studenti: Benjamin Ramovic i Mirnesa Calakovic")
    document.add_paragraph("Novi Pazar, jun 2026.")

    document.add_heading("Uvod i opis problema", level=1)
    document.add_paragraph(
        "Rad obradjuje turisticke recenzije i resava dva prakticna zadatka: multi-label "
        "klasifikaciju oznaka cleanliness, location, luxury i family_friendly, kao i regresionu "
        "procenu sinteticke decimalne ocene posetioca."
    )

    document.add_heading("Opis koriscenih tehnologija i biblioteka", level=1)
    document.add_paragraph(
        "Korisceni su pandas, scikit-learn, sentence-transformers, transformers, PyTorch, "
        "FastAPI i React. Klasicni modeli koriste TF-IDF reprezentaciju, encoder model koristi "
        "sentence-transformers/all-MiniLM-L6-v2, prompt klasifikacija koristi "
        "Qwen/Qwen3-4B-Instruct-2507, a fine-tuning demonstracija koristi prajjwal1/bert-tiny."
    )

    document.add_heading("Opis pocetnog i prosirenog skupa podataka", level=1)
    document.add_paragraph(
        "Pocetni skup sadrzi komentare i cetiri binarne oznake. Dataset je prosiren kolonama "
        "clean_comments i visitor_rating. Visitor_rating je sinteticki target u opsegu 1.0-5.0, "
        "napravljen deterministicki iz labela, tekstualnih signala i kontrolisanog suma."
    )

    document.add_heading("Priprema i obrada podataka", level=1)
    document.add_paragraph(
        "Komentari su ocisceni od HTML tagova, normalizovani su razmaci, uklonjeni su prazni "
        "i duplirani komentari, a label kolone su konvertovane u numericki oblik."
    )
    _add_figure_if_exists(document, figures_path, "label-distribution.png", "Raspodela pozitivnih labela.")
    _add_figure_if_exists(document, figures_path, "review-length-distribution.png", "Raspodela duzina komentara.")
    _add_figure_if_exists(document, figures_path, "visitor-rating-distribution.png", "Raspodela sinteticke ocene.")
    _add_figure_if_exists(document, figures_path, "label-correlation.png", "Korelacije labela.")

    document.add_heading("Klasifikacija i regresija klasicnim modelima", level=1)
    document.add_paragraph(
        "Za klasifikaciju su poredjeni Logistic Regression, Linear SVM i Naive Bayes nad TF-IDF "
        "karakteristikama. Za regresiju su poredjeni Ridge Regression, Linear Regression i Random "
        "Forest Regressor."
    )
    _add_metric_table(
        document,
        "Najbolji klasicni klasifikator",
        [
            ("Model", str(classification.get("best_model", "nije pokrenuto"))),
            ("Micro F1", _metric(classification.get("models", {}).get(classification.get("best_model", ""), {}), "micro_f1")),
            ("Macro F1", _metric(classification.get("models", {}).get(classification.get("best_model", ""), {}), "macro_f1")),
        ],
    )
    _add_metric_table(
        document,
        "Najbolji regresioni model",
        [
            ("Model", str(regression.get("best_model", "nije pokrenuto"))),
            ("MAE", _metric(regression.get("models", {}).get(regression.get("best_model", ""), {}), "mae")),
            ("RMSE", _metric(regression.get("models", {}).get(regression.get("best_model", ""), {}), "rmse")),
            ("R2", _metric(regression.get("models", {}).get(regression.get("best_model", ""), {}), "r2")),
        ],
    )

    document.add_heading("Implementacija encoder modela", level=1)
    document.add_paragraph(
        "Encoder eksperiment koristi sentence-transformers/all-MiniLM-L6-v2 za generisanje "
        "semantickih embeddinga komentara. Nad embedding reprezentacijom trenira se "
        "OneVsRest Logistic Regression klasifikator."
    )
    _add_metric_table(
        document,
        "Encoder rezultati",
        [("Model", str(encoder.get("model", "sentence-transformers/all-MiniLM-L6-v2"))), ("Micro F1", _metric(encoder, "micro_f1"))],
    )

    document.add_heading("Implementacija generativnog modela kroz prompt klasifikaciju", level=1)
    document.add_paragraph(
        "Generativni deo koristi Qwen/Qwen3-4B-Instruct-2507, lokalni causal instruction model. "
        "Model je izabran kao prakticno najjaci ne-gated kandidat za MacBook sa 16 GB RAM-a: "
        "znatno je sposobniji od FLAN-T5 modela, dok su Llama 3.2 i Gemma 3 varijante na Hugging "
        "Face-u gated, a 7B/8B modeli su pogodniji za quantized/GGUF pokretanje. Model dobija "
        "label-by-label prompt sa komentarom i vraca odgovor yes/no koji se parsira u binarnu oznaku."
    )
    _add_metric_table(
        document,
        "Prompt klasifikacija",
        [("Model", str(prompt.get("model", "Qwen/Qwen3-4B-Instruct-2507"))), ("Micro F1", _metric(prompt, "micro_f1"))],
    )
    if isinstance(prompt.get("micro_f1"), float) and prompt["micro_f1"] < 0.2:
        document.add_paragraph(
            "Dobijeni rezultat je nizak, sto pokazuje ogranicenje malog zero-shot generativnog "
            "modela na ovom domenski specificnom multi-label zadatku. Eksperiment je ipak koristan "
            "jer prikazuje ceo tok promptovanja, parsiranja strukturiranog odgovora i evaluacije "
            "u odnosu na ground truth labele."
        )

    document.add_heading("Fine-tuning modela", level=1)
    document.add_paragraph(
        "Fine-tuning je uradjen kao CPU-friendly demonstracija sa prajjwal1/bert-tiny modelom. "
        "Standardni BERT je znatno veci, pa je tiny varijanta izabrana da bi se pokazao tok "
        "fine-tuning-a u lokalnim uslovima."
    )
    _add_metric_table(
        document,
        "Fine-tuning rezultati",
        [("Model", str(finetuning.get("model", "prajjwal1/bert-tiny"))), ("Micro F1", _metric(finetuning, "micro_f1"))],
    )

    document.add_heading("Poredjenje modela i analiza rezultata", level=1)
    document.add_paragraph(
        "Klasicni TF-IDF modeli sluze kao bazna linija. Encoder pristup proverava korist "
        "semantickih reprezentacija, prompt pristup proverava zero/few-shot generativno ponasanje, "
        "a fine-tuning pokazuje kako se transformer moze prilagoditi konkretnom skupu."
    )
    _add_figure_if_exists(document, figures_path, "model-comparison.png", "Poredjenje micro F1 vrednosti.")

    document.add_heading("Analiza gresaka modela", level=1)
    class_errors = error_analysis.get("classification", {})
    reg_errors = error_analysis.get("regression", {})
    document.add_paragraph(
        f"Analiza gresaka sadrzi false positive i false negative primere za {len(class_errors)} labela, "
        f"kao i regresione primere sa najvecim apsolutnim greskama: "
        f"{len(reg_errors.get('largest_absolute_errors', []))} primera."
    )

    document.add_heading("Razvoj jednostavne aplikacije za prikaz rezultata", level=1)
    document.add_paragraph(
        "Aplikacija je podeljena na FastAPI backend i React frontend. Backend ucitava sacuvane "
        "klasicne modele i izlaganjem /predict endpointa vraca label predikcije, verovatnoce "
        "i visitor_rating. Frontend salje komentar i prikazuje rezultate korisniku."
    )

    document.add_heading("Zakljucak", level=1)
    document.add_paragraph(
        "Projekat pokriva ceo tok od pripreme podataka i vizualizacije do klasicnih modela, "
        "encoder pristupa, generativne prompt klasifikacije, fine-tuning demonstracije, "
        "analize gresaka i jednostavne aplikacije za inferencu."
    )

    document.add_heading("Literatura", level=1)
    document.add_paragraph("scikit-learn dokumentacija")
    document.add_paragraph("Hugging Face transformers dokumentacija")
    document.add_paragraph("Sentence Transformers dokumentacija")
    document.add_paragraph("FastAPI i React dokumentacija")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def main() -> None:
    path = generate_report()
    print(f"Wrote report to {path}")


if __name__ == "__main__":
    main()
